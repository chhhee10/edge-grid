"""Assemble the VTU Phase-1 final report as a .docx from the Markdown chapters.

The department submits Word documents, and the chapters are written in Markdown
so they can be diffed and reviewed. This script is the bridge. It applies the
VTU formatting conventions - Times New Roman 12 pt, 1.5 line spacing, justified
body, numbered headings, figures centred with numbered captions - so that the
formatting is a property of the build and not of whoever last edited the file.

    python docs/build_report.py                  # -> docs/report/EdgeGrid_Phase1_Report.docx
    python docs/build_report.py --open           # and open it
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

HERE = Path(__file__).resolve().parent
REPO = HERE.parent
REPORT_DIR = HERE / "report"
FIG_DIR = HERE / "figures"

BODY_FONT = "Times New Roman"
MONO_FONT = "Consolas"

TITLE = "DePIN-Edge: A Decentralized Physical Infrastructure Network for Localized, Verifiable AI Inference"
SUBTITLE = "The Edge Grid"
TEAM = [
    ("HARSHIT RAJ", "1MV23IC021"),
    ("CHETAN RAGHUVANSHI", "1MV23IC013"),
    ("KESHAV NARAYAN", "1MV23IC023"),
    ("MAYUR AGARWAL", "1MV23IC028"),
]
GUIDE = "Dr. SAVITA CHOUDHARY"
GUIDE_TITLE = "Professor & Head"
DEPT = "Department of Computer Science and Engineering\n(IoT, Cyber-Security and Blockchain Technology)"
COLLEGE = "SIR M. VISVESVARAYA INSTITUTE OF TECHNOLOGY"
COLLEGE_ADDR = ("Krishnadevaraya Nagar, International Airport Road,\n"
                "Hunasmaranahalli, Bengaluru - 562157")
UNIVERSITY = "VISVESVARAYA TECHNOLOGICAL UNIVERSITY, BELAGAVI"
UNIVERSITY_ADDR = "Jnana Sangama, Belagavi - 590 018"
YEAR = "2026 - 2027"

# chapters in submission order
CHAPTER_FILES = [
    "ch1_3_introduction.md",
    "ch4_6_literature.md",
    "ch7_proposed_system.md",
    "ch8_results.md",          # written once the experiments have run
    "ch9_conclusion.md",
]


# ----------------------------------------------------------------- helpers

def _style(doc: Document) -> None:
    n = doc.styles["Normal"]
    n.font.name = BODY_FONT
    n.font.size = Pt(12)
    n.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    pf = n.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, bold, before, after in (
        ("Heading 1", 16, True, 18, 12),
        ("Heading 2", 14, True, 14, 8),
        ("Heading 3", 12.5, True, 12, 6),
        ("Heading 4", 12, True, 10, 6),
    ):
        s = doc.styles[name]
        s.font.name = BODY_FONT
        s.font.size = Pt(size)
        s.font.bold = bold
        s.font.color.rgb = RGBColor(0, 0, 0)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        s.paragraph_format.keep_with_next = True


def _centred(doc, text, size=12, bold=False, space_after=6, caps=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    for i, line in enumerate(text.split("\n")):
        r = p.add_run(("\n" if i else "") + (line.upper() if caps else line))
        r.font.size = Pt(size)
        r.bold = bold
        r.font.name = BODY_FONT
    return p


def _page_number_footer(section) -> None:
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


def cover(doc: Document) -> None:
    _centred(doc, UNIVERSITY, 14, True, 2)
    _centred(doc, UNIVERSITY_ADDR, 12, False, 18)
    _centred(doc, "PROJECT PHASE-1 REPORT", 15, True, 18)
    _centred(doc, TITLE, 15, True, 4)
    _centred(doc, f'"{SUBTITLE}"', 13, True, 18)
    _centred(doc,
             "Submitted in partial fulfillment of the requirements for the Sixth Semester\n"
             "BACHELOR OF ENGINEERING\nin\nCOMPUTER SCIENCE AND ENGINEERING\n"
             "(IoT, Cyber-Security and Blockchain Technology)\n"
             f"for the Academic Year {YEAR}", 12, False, 18)
    _centred(doc, "Submitted by", 12, True, 8)
    t = doc.add_table(rows=len(TEAM), cols=2)
    t.alignment = 1
    for row, (name, usn) in zip(t.rows, TEAM):
        for cell, val, bold in ((row.cells[0], name, True), (row.cells[1], usn, False)):
            cp = cell.paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(2)
            r = cp.add_run(val); r.bold = bold; r.font.size = Pt(12); r.font.name = BODY_FONT
    doc.add_paragraph()
    _centred(doc, "Under the guidance of", 12, False, 4)
    _centred(doc, GUIDE, 13, True, 2)
    _centred(doc, GUIDE_TITLE, 12, False, 18)
    _centred(doc, DEPT, 12, True, 6)
    _centred(doc, COLLEGE, 13, True, 2)
    _centred(doc, COLLEGE_ADDR, 11, False, 0)
    doc.add_page_break()


def toc(doc: Document) -> None:
    """A real Word TOC field - Word fills it on F9 / print preview."""
    doc.add_heading("TABLE OF CONTENTS", level=1)
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-3" \h \z \u')
    run = OxmlElement("w:r"); t = OxmlElement("w:t")
    t.text = "Right-click and choose 'Update Field' to build the table of contents."
    run.append(t); fld.append(run); p._p.append(fld)
    doc.add_page_break()


# ------------------------------------------------------------- md -> docx

INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`|\[[0-9a-z]+\])")


def _runs(par, text: str) -> None:
    """Bold, italic, inline code and citation markers, applied per run."""
    for piece in INLINE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            par.add_run(piece[2:-2]).bold = True
        elif piece.startswith("*") and piece.endswith("*") and len(piece) > 2:
            par.add_run(piece[1:-1]).italic = True
        elif piece.startswith("`") and piece.endswith("`"):
            r = par.add_run(piece[1:-1]); r.font.name = MONO_FONT; r.font.size = Pt(10.5)
        else:
            par.add_run(piece)


def _table(doc, rows: list[str]) -> None:
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    cells = [c for c in cells if not all(set(x) <= set("-: ") for x in c)]
    if not cells:
        return
    t = doc.add_table(rows=len(cells), cols=len(cells[0]))
    t.style = "Table Grid"
    for i, row in enumerate(cells):
        for j, val in enumerate(row[: len(cells[0])]):
            p = t.rows[i].cells[j].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.0
            _runs(p, val)
            for r in p.runs:
                r.font.size = Pt(10.5); r.font.name = BODY_FONT
                if i == 0:
                    r.bold = True
    doc.add_paragraph()


def render(doc: Document, md: str, fig_counter: list[int]) -> None:
    lines = md.split("\n")
    i, buf = 0, []

    def flush():
        nonlocal buf
        if buf:
            text = " ".join(x.strip() for x in buf).strip()
            if text:
                _runs(doc.add_paragraph(), text)
            buf = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            flush()
            i += 1
            code = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i]); i += 1
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.left_indent = Inches(0.3)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run("\n".join(code)); r.font.name = MONO_FONT; r.font.size = Pt(9.5)
            i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            flush()
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(lines[i]); i += 1
            _table(doc, rows)
            continue

        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            flush()
            level, text = len(m.group(1)), m.group(2).strip()
            if level == 1 and re.match(r"^Chapter\s+\d", text, re.I):
                doc.add_page_break()
            doc.add_heading(text, level=min(level, 4))
            i += 1
            continue

        # Two figure conventions are accepted: standard Markdown images, and the
        # chapters' own "**Figure N.N** - `path` - caption" caption lines.
        m = re.match(r"^!\[(.*?)\]\((.*?)\)$", stripped)
        fig_line = None
        if not m:
            fig_line = re.match(
                r"^\*\*(Figure\s+[\d.]+)\*\*\s*[-\u2014]+\s*`([^`]*docs/figures/[^`]+)`\s*[-\u2014]+\s*(.*)$",
                stripped)
        if fig_line:
            flush()
            label, src, cap = fig_line.group(1), fig_line.group(2), fig_line.group(3)
            # a caption may wrap onto following lines until a blank line
            j = i + 1
            while j < len(lines) and lines[j].strip() and not lines[j].strip().startswith(("#", "|", "**Figure")):
                cap += " " + lines[j].strip(); j += 1
            path = REPO / src
            if path.suffix == ".svg":
                path = path.with_suffix(".png")
            if path.exists():
                doc.add_picture(str(path), width=Inches(6.1))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                fig_counter[0] += 1
                c = doc.add_paragraph()
                c.alignment = WD_ALIGN_PARAGRAPH.CENTER
                c.paragraph_format.space_after = Pt(12)
                r = c.add_run(f"{label}: {cap.rstrip('.')}")
                r.font.size = Pt(10.5); r.italic = True
            else:
                _runs(doc.add_paragraph(), f"[missing figure: {src}]")
            i = j
            continue

        if m:
            flush()
            caption_text, src = m.group(1), m.group(2)
            path = (REPO / src) if not src.startswith("/") else Path(src)
            if path.suffix == ".svg":
                path = path.with_suffix(".png")
            if path.exists():
                doc.add_picture(str(path), width=Inches(6.0))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                fig_counter[0] += 1
                c = doc.add_paragraph()
                c.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = c.add_run(f"Fig. {fig_counter[0]}: {caption_text}")
                r.font.size = Pt(10.5); r.italic = True
            else:
                _runs(doc.add_paragraph(), f"[missing figure: {src}]")
            i += 1
            continue

        m = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)$", line)
        if m:
            flush()
            indent, marker, text = m.group(1), m.group(2), m.group(3)
            style = "List Number" if marker[0].isdigit() else "List Bullet"
            p = doc.add_paragraph(style=style)
            p.paragraph_format.left_indent = Inches(0.35 + 0.3 * (len(indent) // 2))
            p.paragraph_format.space_after = Pt(3)
            _runs(p, text)
            i += 1
            continue

        if not stripped or set(stripped) <= set("-*_") and len(stripped) >= 3:
            flush()
            i += 1
            continue

        buf.append(line)
        i += 1

    flush()


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--out", type=Path,
                    default=REPORT_DIR / "EdgeGrid_Phase1_Report.docx")
    ap.add_argument("--chapters", nargs="*", default=None)
    args = ap.parse_args(argv)

    doc = Document()
    for s in doc.sections:
        # A4 with a wider left margin for binding - the VTU convention.
        s.page_width, s.page_height = Mm(210), Mm(297)
        s.top_margin = s.bottom_margin = Mm(25)
        s.left_margin = Mm(32); s.right_margin = Mm(25)
    _style(doc)
    cover(doc)
    toc(doc)
    _page_number_footer(doc.sections[0])

    fig_counter = [0]
    files = args.chapters or CHAPTER_FILES
    included, missing = [], []
    for name in files:
        path = REPORT_DIR / name
        if not path.exists():
            missing.append(name)
            continue
        render(doc, path.read_text(encoding="utf-8"), fig_counter)
        included.append(name)

    refs = HERE / "REFERENCES.md"
    if refs.exists():
        doc.add_page_break()
        render(doc, refs.read_text(encoding="utf-8"), fig_counter)
        included.append("REFERENCES.md")

    args.out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.out)

    words = sum(len((REPORT_DIR / f).read_text().split()) for f in included
                if (REPORT_DIR / f).exists())
    print(f"wrote {args.out}")
    print(f"  chapters : {', '.join(included)}")
    print(f"  figures  : {fig_counter[0]}")
    print(f"  ~words   : {words:,}")
    if missing:
        print(f"  PENDING  : {', '.join(missing)} (not yet written - run the experiments first)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
